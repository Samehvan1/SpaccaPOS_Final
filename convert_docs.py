import os
import re
import subprocess
import markdown
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Deep Navy
        # Add bottom border or horizontal line effect
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3) # Primary Blue
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGBColor(0x2B, 0x2D, 0x42) # Slate Dark
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    return p

def format_inline_markdown(paragraph, text):
    # Regex to split by bold (**text**), italic (*text*), code (`text`), math ($...$)
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\$.*?\$)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
            r.font.name = 'Calibri'
        elif token.startswith('*') and token.endswith('*'):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
            r.font.name = 'Calibri'
        elif token.startswith('`') and token.endswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            r = paragraph.add_run(token)
            r.font.name = 'Calibri'

def convert_md_to_docx(md_filepath, docx_filepath):
    doc = Document()
    
    # Page Margins (1 inch everywhere)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    with open(md_filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []
    in_alert = False
    alert_type = 'NOTE'
    alert_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Code block handler
        if stripped.startswith('```'):
            if in_code_block:
                # Flush code block
                in_code_block = False
                code_text = "".join(code_block_lines)
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                set_cell_background(cell, "F4F6F8")
                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(code_text)
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
                code_block_lines = []
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue
            
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
            
        # Table handler
        if stripped.startswith('|') and stripped.endswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            in_table = False
            # Process table_lines
            rows_data = []
            for tline in table_lines:
                # Ignore separator lines like |---|---|
                if re.match(r'^\|[\s\:\-\|]+\|$', tline):
                    continue
                cells = [c.strip() for c in tline.split('|')[1:-1]]
                if cells:
                    rows_data.append(cells)
            
            if rows_data:
                cols_cnt = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=cols_cnt)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for row_idx, row_content in enumerate(rows_data):
                    for col_idx, cell_value in enumerate(row_content):
                        if col_idx < cols_cnt:
                            cell = table.cell(row_idx, col_idx)
                            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(2)
                            p.paragraph_format.space_after = Pt(2)
                            
                            if row_idx == 0:
                                set_cell_background(cell, "1B365D") # Navy Header
                                format_inline_markdown(p, cell_value)
                                for r in p.runs:
                                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    r.bold = True
                            else:
                                if row_idx % 2 == 1:
                                    set_cell_background(cell, "FFFFFF")
                                else:
                                    set_cell_background(cell, "F7FAFC") # Light zebra tint
                                format_inline_markdown(p, cell_value)
            table_lines = []
            
        # Alert / Blockquote handler
        if stripped.startswith('>'):
            alert_text = stripped[1:].strip()
            if alert_text.startswith('[!') and ']' in alert_text:
                alert_type = alert_text[2:alert_text.find(']')].upper()
                alert_text = alert_text[alert_text.find(']')+1:].strip()
            
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            
            bg_color = "EDF2F7"
            border_color = "4A5568"
            if alert_type == 'NOTE':
                bg_color = "EBF8FF"
                border_color = "3182CE"
            elif alert_type in ('CAUTION', 'WARNING'):
                bg_color = "FFF5F5"
                border_color = "E53E3E"
            elif alert_type == 'IMPORTANT':
                bg_color = "FEFCBF"
                border_color = "D69E2E"
                
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            r_head = p.add_run(f"[{alert_type}] ")
            r_head.bold = True
            r_head.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0) if alert_type=='NOTE' else RGBColor(0xC5, 0x30, 0x30)
            format_inline_markdown(p, alert_text)
            i += 1
            continue
            
        # Headings
        if stripped.startswith('#'):
            h_match = re.match(r'^(#+)\s+(.*)$', stripped)
            if h_match:
                level = len(h_match.group(1))
                text = h_match.group(2)
                add_styled_heading(doc, text, level)
                i += 1
                continue
                
        # Horizontal Rule
        if stripped in ('---', '***', '___'):
            # Add subtle spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
            
        # Bullet list / Numbered list
        if re.match(r'^[\*\-\+]\s+', stripped):
            list_text = re.sub(r'^[\*\-\+]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            format_inline_markdown(p, list_text)
            i += 1
            continue
            
        if re.match(r'^\d+\.\s+', stripped):
            list_text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            format_inline_markdown(p, list_text)
            i += 1
            continue
            
        # Normal Paragraph
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            format_inline_markdown(p, stripped)
            
        i += 1
        
    doc.save(docx_filepath)
    print(f"Generated DOCX: {docx_filepath}")

def convert_md_to_pdf(md_filepath, html_filepath, pdf_filepath):
    with open(md_filepath, 'r', encoding='utf-8') as f:
        md_content = f.read()
        
    # Convert Mermaid codeblocks to pre formatted styled text for clean printing
    md_content = re.sub(r'```mermaid\s([\s\S]*?)```', r'<pre class="mermaid-diagram">\1</pre>', md_content)
    
    html_body = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'toc'])
    
    full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>SpaccaPOS Documentation</title>
<style>
    @page {{
        size: A4;
        margin: 20mm 15mm 20mm 15mm;
        @bottom-right {{
            content: counter(page);
        }}
    }}
    body {{
        font-family: 'Segoe UI', Arial, sans-serif;
        color: #2D3748;
        line-height: 1.6;
        font-size: 11pt;
        background-color: #FFFFFF;
    }}
    h1 {{
        color: #1B365D;
        font-size: 22pt;
        border-bottom: 2px solid #1B365D;
        padding-bottom: 6px;
        margin-top: 24pt;
        margin-bottom: 12pt;
        page-break-before: always;
    }}
    h1:first-of-type {{
        page-break-before: avoid;
    }}
    h2 {{
        color: #0056B3;
        font-size: 16pt;
        border-bottom: 1px solid #E2E8F0;
        padding-bottom: 4px;
        margin-top: 18pt;
        margin-bottom: 8pt;
    }}
    h3 {{
        color: #2B2D42;
        font-size: 13pt;
        margin-top: 14pt;
        margin-bottom: 6pt;
    }}
    h4 {{
        color: #4A5568;
        font-size: 11pt;
        margin-top: 10pt;
        margin-bottom: 4pt;
    }}
    p {{
        margin-top: 0;
        margin-bottom: 8pt;
    }}
    table {{
        width: 100%;
        border-collapse: collapse;
        margin-top: 12pt;
        margin-bottom: 16pt;
        font-size: 10pt;
        page-break-inside: avoid;
    }}
    th {{
        background-color: #1B365D;
        color: #FFFFFF;
        font-weight: bold;
        text-align: left;
        padding: 8px 10px;
        border: 1px solid #1B365D;
    }}
    td {{
        padding: 8px 10px;
        border: 1px solid #CBD5E0;
    }}
    tr:nth-child(even) td {{
        background-color: #F7FAFC;
    }}
    blockquote {{
        border-left: 4px solid #3182CE;
        background-color: #EBF8FF;
        margin: 12pt 0;
        padding: 10pt 14pt;
        border-radius: 4px;
    }}
    pre, code {{
        font-family: 'Consolas', 'Courier New', monospace;
    }}
    code {{
        background-color: #EDF2F7;
        color: #C7254E;
        padding: 2px 5px;
        border-radius: 3px;
        font-size: 9.5pt;
    }}
    pre {{
        background-color: #2D3748;
        color: #F7FAFC;
        padding: 12px;
        border-radius: 6px;
        overflow-x: auto;
        font-size: 9pt;
        line-height: 1.4;
    }}
    pre code {{
        background-color: transparent;
        color: inherit;
        padding: 0;
    }}
    .mermaid-diagram {{
        background-color: #F7FAFC;
        color: #1A202C;
        border: 1px solid #E2E8F0;
        border-left: 4px solid #0056B3;
        font-family: 'Consolas', monospace;
        white-space: pre-wrap;
    }}
    ul, ol {{
        margin-top: 0;
        margin-bottom: 8pt;
        padding-left: 20pt;
    }}
    li {{
        margin-bottom: 4pt;
    }}
    hr {{
        border: 0;
        height: 1px;
        background: #E2E8F0;
        margin: 18pt 0;
    }}
</style>
</head>
<body>
{html_body}
</body>
</html>
"""
    with open(html_filepath, 'w', encoding='utf-8') as f:
        f.write(full_html)
        
    # Execute MS Edge headless to generate PDF
    edge_path = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
    cmd = [
        edge_path,
        "--headless",
        "--disable-gpu",
        f"--print-to-pdf={pdf_filepath}",
        "--no-margins",
        os.path.abspath(html_filepath)
    ]
    subprocess.run(cmd, check=True)
    print(f"Generated PDF: {pdf_filepath}")

if __name__ == '__main__':
    base_dir = r"d:\MyWorks\SpaccaTests\SpaccaPos20260416_0\SpaccaPos"
    
    # 1. SpaccaPOS SRS
    srs_md = os.path.join(base_dir, "SpaccaPOS_SRS.md")
    srs_docx = os.path.join(base_dir, "SpaccaPOS_SRS.docx")
    srs_html = os.path.join(base_dir, "SpaccaPOS_SRS.html")
    srs_pdf = os.path.join(base_dir, "SpaccaPOS_SRS.pdf")
    
    convert_md_to_docx(srs_md, srs_docx)
    convert_md_to_pdf(srs_md, srs_html, srs_pdf)
    
    # 2. SpaccaPOS User & Staff Manual
    manual_md = os.path.join(base_dir, "SpaccaPOS_User_and_Staff_Manual.md")
    manual_docx = os.path.join(base_dir, "SpaccaPOS_User_and_Staff_Manual.docx")
    manual_html = os.path.join(base_dir, "SpaccaPOS_User_and_Staff_Manual.html")
    manual_pdf = os.path.join(base_dir, "SpaccaPOS_User_and_Staff_Manual.pdf")
    
    convert_md_to_docx(manual_md, manual_docx)
    convert_md_to_pdf(manual_md, manual_html, manual_pdf)
    
    print("SUCCESS: All documents successfully converted to Word (.docx) and PDF (.pdf)!")
